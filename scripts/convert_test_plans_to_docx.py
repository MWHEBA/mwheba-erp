#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Universal Markdown to 100% Strict Root-RTL Word (.docx) Converter for MWHEBA ERP QA Plans
محول ملفات Markdown لـ Word بتنسيق عربي RTL 100%، محاذاة يمين كاملة، ومعالجة أخطاء قفل الملفات عند فتحها في Word.
"""

import os
import sys
import re
from pathlib import Path

# ضبط ترميز الطرفية في Windows
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_NAME = "Cairo"


def apply_global_arabic_rtl_root(doc):
    """تطبيق التوجيه العربي ومحاذاة اليمين على الجذور والطبقات الست للمستند"""
    
    # 1. طبقة إعدادات المستند (Document Settings)
    settings_elem = doc.settings.element
    theme_lang = parse_xml(f'<w:themeFontLang {nsdecls("w")} w:val="en-US" w:bidi="ar-EG"/>')
    settings_elem.append(theme_lang)
    
    # 2. طبقة الافتراضيات العامة (docDefaults in styles.xml)
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        pPrDefault = doc_defaults.find(qn('w:pPrDefault'))
        if pPrDefault is None:
            pPrDefault = parse_xml(f'<w:pPrDefault {nsdecls("w")}><w:pPr><w:bidi w:val="1"/><w:jc w:val="right"/></w:pPr></w:pPrDefault>')
            doc_defaults.append(pPrDefault)
        else:
            pPr = pPrDefault.find(qn('w:pPr'))
            if pPr is None:
                pPrDefault.append(parse_xml(f'<w:pPr {nsdecls("w")}><w:bidi w:val="1"/><w:jc w:val="right"/></w:pPr>'))
            else:
                pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
                pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="right"/>'))
                
        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is not None:
                rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'))
                rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>'))
                rPr.append(parse_xml(f'<w:lang {nsdecls("w")} w:val="ar-EG" w:bidi="ar-EG"/>'))

    # 3. طبقة النمط الافتراضي (Normal Style)
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(9.5)
    pPr = style.element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="right"/>'))
    rPr = style.element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>'))
    rPr.append(parse_xml(f'<w:lang {nsdecls("w")} w:val="ar-EG" w:bidi="ar-EG"/>'))
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>'))

    # 4. طبقة الأقسام (Sections)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        sectPr = section._sectPr
        sectPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))
        sectPr.append(parse_xml(f'<w:rtlGutter {nsdecls("w")} w:val="1"/>'))


def build_strict_pPr(shd_hex=None, bdr_color=None, bdr_sz=None, before_pt=0, after_pt=0, jc_val="right"):
    """بناء كائن pPr وفقاً للترتيب الصارم لمخطط OpenXML مع محاذاة يمين صارمة"""
    parts = [f'<w:pPr {nsdecls("w")}>']
    
    if bdr_color:
        sz_val = bdr_sz if bdr_sz else "6"
        parts.append(f'<w:pBdr><w:bottom w:val="single" w:sz="{sz_val}" w:space="4" w:color="{bdr_color}"/></w:pBdr>')
        
    if shd_hex:
        parts.append(f'<w:shd w:fill="{shd_hex}"/>')
        
    parts.append('<w:bidi w:val="1"/>')
    
    if before_pt > 0 or after_pt > 0:
        b_val = int(before_pt * 20)
        a_val = int(after_pt * 20)
        parts.append(f'<w:spacing w:before="{b_val}" w:after="{a_val}"/>')
        
    parts.append(f'<w:jc w:val="{jc_val}"/>')
    parts.append('</w:pPr>')
    return parse_xml(''.join(parts))


def apply_pPr_to_paragraph(p, shd_hex=None, bdr_color=None, bdr_sz=None, before_pt=0, after_pt=0, jc_val="right"):
    """تطبيق خصائص الفقرة مع استبدال pPr القديم لضمان صحة ترتيب العناصر"""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    old_pPr = p._element.find(qn('w:pPr'))
    if old_pPr is not None:
        p._element.remove(old_pPr)
        
    new_pPr = build_strict_pPr(
        shd_hex=shd_hex,
        bdr_color=bdr_color,
        bdr_sz=bdr_sz,
        before_pt=before_pt,
        after_pt=after_pt,
        jc_val=jc_val
    )
    p._element.insert(0, new_pPr)


def set_cell_background(cell, fill_hex):
    """تعيين لون خلفية خلية الجدول"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """ضبط الهوامش الداخلية للخلايا"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """تطبيق حدود أنيقة وناعمة على الجدول بالكامل"""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="{val}" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="{val}" w:sz="4" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)


def add_styled_text(paragraph, text, is_bold=False, is_italic=False, color_rgb=None, font_size_pt=9.5, font_family=FONT_NAME):
    """إضافة نصوص منسقة مع تفعيل حجم الخط العربي واللاتيني ومحاذاة اليمين"""
    text = text.replace("[ ]", "☐ ").replace("[x]", "☑ ").replace("[X]", "☑ ")
    
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        
        token_bold = is_bold
        token_italic = is_italic
        token_color = color_rgb
        
        if token.startswith("**") and token.endswith("**"):
            token_text = token[2:-2]
            token_bold = True
            if not token_color:
                token_color = RGBColor(15, 23, 42)
        elif token.startswith("*") and token.endswith("*"):
            token_text = token[1:-1]
            token_italic = True
        elif token.startswith("`") and token.endswith("`"):
            token_text = token[1:-1]
            token_color = RGBColor(185, 28, 28)
            token_bold = True
        else:
            token_text = token
            
        run.text = token_text
        
        sz_val = int(font_size_pt * 2)
        
        rPr_parts = [
            f'<w:rPr {nsdecls("w")}>',
            f'<w:rFonts w:ascii="{font_family}" w:hAnsi="{font_family}" w:cs="{font_family}"/>',
            f'<w:sz w:val="{sz_val}"/>',
            f'<w:szCs w:val="{sz_val}"/>',
            '<w:rtl w:val="1"/>',
            '<w:lang w:val="ar-EG" w:bidi="ar-EG"/>',
        ]
        
        if token_bold:
            rPr_parts.append('<w:b/><w:bCs/>')
        if token_italic:
            rPr_parts.append('<w:i/><w:iCs/>')
        if token_color:
            color_hex = f"{token_color[0]:02X}{token_color[1]:02X}{token_color[2]:02X}"
            rPr_parts.append(f'<w:color w:val="{color_hex}"/>')
            
        rPr_parts.append('</w:rPr>')
        
        rPr = parse_xml(''.join(rPr_parts))
        run._element.insert(0, rPr)


def convert_markdown_to_docx(md_file_path, output_docx_path=None):
    """تحويل ملف Markdown إلى Word بمحاذاة يمين صارمة 100% وأحجام هادئة واحترافية"""
    md_file = Path(md_file_path)
    if not md_file.exists():
        print(f"[ERROR] File not found: {md_file_path}")
        return False
        
    if not output_docx_path:
        output_docx_path = md_file.with_suffix(".docx")
    else:
        output_docx_path = Path(output_docx_path)
        
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = docx.Document()

    # تطبيق الجذر الشامل لـ RTL والمحاذاة اليمينية
    apply_global_arabic_rtl_root(doc)

    lines = md_content.splitlines()
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        # معالجة كتل الكود (Code blocks)
        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                apply_pPr_to_paragraph(p, shd_hex="F8FAFC", before_pt=3, after_pt=3, jc_val="right")
                add_styled_text(p, "\n".join(code_lines), font_size_pt=8.5, color_rgb=RGBColor(30, 41, 59))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # معالجة الجداول
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            cols = [col.strip() for col in stripped[1:-1].split("|")]
            table_rows.append(cols)
            in_table = True
            continue
        elif in_table:
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table, color="CBD5E1", sz="4")
                
                tblPr = table._element.xpath('w:tblPr')
                if tblPr:
                    tblBidi = parse_xml(f'<w:bidiVisual {nsdecls("w")} w:val="1"/>')
                    tblJc = parse_xml(f'<w:jc {nsdecls("w")} w:val="right"/>')
                    tblPr[0].append(tblBidi)
                    tblPr[0].append(tblJc)

                for r_idx, row_data in enumerate(table_rows):
                    row = table.rows[r_idx]
                    is_header = (r_idx == 0)
                    
                    bg_color = "1E293B" if is_header else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
                    text_color = RGBColor(255, 255, 255) if is_header else RGBColor(15, 23, 42)

                    for c_idx in range(num_cols):
                        cell = row.cells[c_idx]
                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                        p = cell.paragraphs[0]
                        apply_pPr_to_paragraph(p, before_pt=1, after_pt=1, jc_val="right")
                        add_styled_text(p, cell_text, is_bold=is_header, color_rgb=text_color, font_size_pt=8.5 if not is_header else 9)

                doc.add_paragraph()
            in_table = False
            table_rows = []

        # تجاهل الفواصل الأفقية
        if stripped in ["---", "***", "___"]:
            continue

        if not stripped:
            continue

        # =========================================================================
        # 👑 تدرج العناوين الهادئ والمحاذاة اليمينية الصارمة
        # =========================================================================

        # 👑 H1 - العنوان الرئيسي للمستند (15 pt عريض، كحلي، خط سفلي خفيف)
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, bdr_color="0284C7", bdr_sz="8", before_pt=12, after_pt=4, jc_val="right")
            add_styled_text(p, stripped[2:], is_bold=True, color_rgb=RGBColor(15, 23, 42), font_size_pt=15)

        # 📌 H2 - العناوين الفرعية الكبرى / الفهرس (13 pt عريض، رمادي داكن)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=10, after_pt=3, jc_val="right")
            add_styled_text(p, stripped[3:], is_bold=True, color_rgb=RGBColor(30, 41, 59), font_size_pt=13)

        # 🏷️ H3 - عناوين الأقسام الرئيسية الـ 11 (11.5 pt عريض، أزرق كحلي داكن)
        elif stripped.startswith("### "):
            heading_text = stripped[4:]
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, bdr_color="E2E8F0", bdr_sz="4", before_pt=10, after_pt=3, jc_val="right")
            add_styled_text(p, heading_text, is_bold=True, color_rgb=RGBColor(15, 23, 42), font_size_pt=11.5)

        # 🔍 H4 - بطاقات حالات الاختبار TC-xxx (10.5 pt عريض، أزرق مميز)
        elif stripped.startswith("#### "):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=6, after_pt=2, jc_val="right")
            add_styled_text(p, stripped[5:], is_bold=True, color_rgb=RGBColor(2, 132, 199), font_size_pt=10.5)

        # 💬 الاقتباسات والملاحظات الإرشادية (> Note)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, shd_hex="F8FAFC", before_pt=2, after_pt=2, jc_val="right")
            add_styled_text(p, stripped[2:], is_italic=True, color_rgb=RGBColor(71, 85, 105), font_size_pt=9)

        # 🔲 بنود الاختيار والفحص (Checklists: - [ ])
        elif stripped.startswith("- [ ]") or stripped.startswith("- [x]") or stripped.startswith("- [X]"):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=1, after_pt=1, jc_val="right")
            add_styled_text(p, stripped[2:], font_size_pt=9.5, color_rgb=RGBColor(30, 41, 59))

        # 🔹 القوائم النقطية والرقمية العادية
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=1, after_pt=1, jc_val="right")
            add_styled_text(p, "▪ " + stripped[2:], font_size_pt=9.5)

        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=1, after_pt=1, jc_val="right")
            add_styled_text(p, stripped, font_size_pt=9.5)

        # 📝 الفقرات النصية العادية
        else:
            p = doc.add_paragraph()
            apply_pPr_to_paragraph(p, before_pt=2, after_pt=2, jc_val="right")
            add_styled_text(p, stripped, font_size_pt=9.5, color_rgb=RGBColor(51, 65, 85))

    try:
        doc.save(str(output_docx_path))
        print(f"[SUCCESS] Converted to 100% Root-RTL Flush-Right Word (Cairo Font): {output_docx_path}")
        return True
    except PermissionError:
        print(f"[WARN] Could not save '{output_docx_path}' because it is open in Microsoft Word. Please close Word and run again.")
        return False


def convert_all_test_plans():
    """تحويل كافة ملفات الـ Markdown الموجودة في مجلد خطط الاختبار إلى Word"""
    base_dir = Path(__file__).resolve().parent.parent
    qa_dir = base_dir / "docs" / "qa_test_plans"
    
    md_files = []
    if qa_dir.exists():
        md_files.extend(list(qa_dir.glob("*.md")))
        
    root_master = base_dir / "TESTING_MASTER_PLAN.md"
    if root_master.exists():
        md_files.append(root_master)

    print(f"[START] Converting {len(md_files)} Markdown files to 100% Root-RTL Word DOCX...")
    for md_file in md_files:
        docx_path = md_file.with_suffix(".docx")
        convert_markdown_to_docx(md_file, docx_path)

    print("[DONE] Conversion process finished!")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        target_path = sys.argv[1]
        convert_markdown_to_docx(target_path)
    else:
        convert_all_test_plans()
